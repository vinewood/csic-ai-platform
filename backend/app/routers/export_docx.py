"""文档导出路由 — docx 生成（论文格式，仅导出AI回复）"""

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from ..auth import get_current_user
import io, re

router = APIRouter(prefix="/api/export", tags=["导出"])

class ExportRequest(BaseModel):
    conversation_id: int
    format: Optional[str] = "docx"

def clean_markdown(text: str) -> str:
    """去除 Markdown 格式字符，保留纯文本"""
    if not text: return ""
    # 去掉加粗/斜体标记
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # 去掉代码块标记
    text = re.sub(r'```[\s\S]*?```', lambda m: m.group(0).replace('```', '').strip(), text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 保留 ## ### 等标题标记（用于识别标题），去掉列表标记
    text = re.sub(r'^[\s]*[-*+]\s', '• ', text, flags=re.MULTILINE)
    # 去掉水平线
    text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    # 去掉多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

@router.post("/docx")
async def export_docx(req: ExportRequest, current_user: dict = Depends(get_current_user)):
    """导出对话的AI回复为论文格式docx"""
    from ..database import async_session
    from ..models import Message, Conversation
    from sqlalchemy import select
    
    async with async_session() as db:
        conversation = await db.get(Conversation, req.conversation_id)
        if not conversation:
            return {"error": "对话不存在"}
        
        result = await db.execute(
            select(Message).where(
                Message.conversation_id == req.conversation_id,
                Message.role == "assistant"
            ).order_by(Message.id)
        )
        messages = result.scalars().all()
    
    if not messages:
        return {"error": "没有AI回复内容"}
    
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title_text = (conversation.title or "AI回复内容").replace('*', '')
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 副标题信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"生成时间: {conversation.created_at}")
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph("")
    
    # 逐条AI回复，解析Markdown结构
    for msg in messages:
        content = clean_markdown(msg.content or "")
        if not content.strip():
            continue
        
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # 检测标题 ## 或 ###
            h_match = re.match(r'^(#{1,3})\s+(.+)', line)
            if h_match:
                level = len(h_match.group(1))
                heading = doc.add_heading(h_match.group(2), level=min(level, 3))
                for run in heading.runs:
                    run.font.name = '黑体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                i += 1
                continue
            
            # 检测编号列表
            num_match = re.match(r'^(\d+)[\.\)、]\s*(.+)', line)
            if num_match:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.space_after = Pt(3)
                p.style.font.size = Pt(12)
                run = p.add_run(f"{num_match.group(1)}. {num_match.group(2)}")
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                i += 1
                continue
            
            # 检测项目符号
            bullet_match = re.match(r'^(•|-)\s+(.+)', line)
            if bullet_match:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"• {bullet_match.group(2)}")
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                i += 1
                continue
            
            # 普通段落：合并连续文本行
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,3}\s|\d+[\.\)、]|•|-)', lines[i]):
                para_lines.append(lines[i].strip())
                i += 1
            
            if para_lines:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
                text = ' '.join(para_lines)
                run = p.add_run(text)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            else:
                i += 1
        
        # 每条消息后加分隔
        doc.add_paragraph("")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{conversation.title or 'AI回复'}.docx"
    from urllib.parse import quote
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )
